"""
Bouwt het uitgebreide Verborgen_Kennis_Studieboek.docx
met extra uitvinders, blueprints en nabouw-instructies.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── helpers ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def h1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p

def h2(doc, text, color=RGBColor(0x16, 0x21, 0x3E)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = color
    return p

def h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x0F, 0x60, 0x0F)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11) if p.runs else None
    return p

def quote(doc, text, author=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'"{text}"')
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    if author:
        p.add_run(f"  — {author}")
    return p

def blueprint_box(doc, title, content_lines):
    """Render a monospaced blueprint block in a shaded table."""
    p = doc.add_paragraph()
    run = p.add_run(f"📐 BLUEPRINT — {title}")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xD4, 0x6A, 0x00)

    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, 'F5F0E8')
    cell.text = ''
    for line in content_lines:
        p2 = cell.add_paragraph(line)
        run2 = p2.runs[0] if p2.runs else p2.add_run(line)
        run2.font.name = 'Courier New'
        run2.font.size = Pt(8.5)
    doc.add_paragraph()

def materials_box(doc, items):
    p = doc.add_paragraph()
    run = p.add_run("🔧 BENODIGDE MATERIALEN")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x6B, 0x3A)
    for item in items:
        doc.add_paragraph(f"  • {item}", style='List Bullet')

def steps_box(doc, steps):
    p = doc.add_paragraph()
    run = p.add_run("⚙️ STAP-VOOR-STAP NABOUW")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6B)
    for i, step in enumerate(steps, 1):
        p2 = doc.add_paragraph()
        p2.add_run(f"Stap {i}: ").bold = True
        p2.add_run(step)

def divider(doc):
    doc.add_paragraph("─" * 80)

# ─── DOCUMENT OPBOUWEN ─────────────────────────────────────────────────────

doc = Document()

# Paginamarges instellen
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ══════════════════════ TITELPAGINA ═══════════════════════════════════════
h1(doc, "VERBORGEN KENNIS")
h1(doc, "Onderdrukte Wetenschap, Kunst & Ontdekkingen")
doc.add_paragraph()
body(doc, "Een studieboek voor kritisch denkers")
body(doc, "Aetherion Vectoris · UVRM Art. 19 & 27")
body(doc, "Editie 2 — Uitgebreid met Blueprints & Nabouwinstructies")
doc.add_page_break()

# ══════════════════════ INTRO ══════════════════════════════════════════════
h2(doc, "  INTRO    Waarom dit document bestaat")
body(doc, "Gedurende eeuwen zijn wetenschappers, kunstenaars en denkers wiens werk het gevestigde belang bedreigde systematisch onderdrukt, gemarginaliseerd of vergeten. Dit is geen complottheorie — het is gedocumenteerde geschiedenis.")
doc.add_paragraph()
body(doc, "Dit document is een studieboek. Per figuur leer je: wie ze waren, wat ze ontdekten, hoe het werd onderdrukt, en — het belangrijkste — welke lessen je er zelf uit kunt trekken voor je eigen denken, creëren en handelen. Editie 2 voegt technische blueprints en nabouwinstructies toe zodat de kennis niet alleen bewaard blijft, maar ook herschapen kan worden.")
quote(doc, "All truth passes through three stages. First, it is ridiculed. Second, it is violently opposed. Third, it is accepted as being self-evident.", "Arthur Schopenhauer")
doc.add_page_break()

# ══════════════════════ 01 TESLA ══════════════════════════════════════════
h2(doc, "  01    Nikola Tesla — De Man die de Wereld Gratis Stroom Wilde Geven")
h3(doc, "Wie was hij?")
body(doc, "Nikola Tesla (1856–1943) was een Servisch-Amerikaans uitvinder en elektrotechnisch ingenieur. Hij werkte aanvankelijk voor Edison maar brak met hem toen hij begreep dat wisselstroom (AC) verre superieur was aan gelijkstroom (DC). Tesla won de 'War of Currents' — het stroomstelsel dat de hele wereld vandaag gebruikt is van hem.")
body(doc, "Maar Tesla's visie reikte veel verder dan wisselstroom. Hij geloofde dat energie draadloos en gratis beschikbaar kon zijn voor iedereen op aarde — zonder meter, zonder rekening, zonder monopolie.")

h3(doc, "Zijn belangrijkste ontdekkingen")
for item in [
    "Wisselstroom (AC): De basis van elk stopcontact ter wereld.",
    "Radiocommunicatie: Tesla demonstreerde draadloze signaaloverdracht vóór Marconi.",
    "Wardenclyffe Tower: 57 meter hoge toren om draadloos energie te verzenden.",
    "369-Frequentietheorie: Tesla's obsessie met 3, 6 en 9 als sleutel tot het universum.",
    "Resonantie & aardtrillingen: Mechanische resonantie als destructieve kracht.",
    "Tesla-spoel: Hoge-frequentietransformator — nog altijd in gebruik.",
]:
    doc.add_paragraph(f"  • {item}", style='List Bullet')

h3(doc, "Hoe werd het onderdrukt?")
body(doc, "De financier J.P. Morgan stopte de financiering van Wardenclyffe Tower zodra hij begreep dat het systeem geen meters kon installeren. Na Tesla's dood in 1943 werden zijn papieren binnen 24 uur in beslag genomen door de US Office of Alien Property. Het FBI-dossier werd pas in de jaren 2000 gedeeltelijk vrijgegeven via FOIA-verzoeken.")

divider(doc)

# TESLA BLUEPRINT 1 — Tesla Spoel
blueprint_box(doc, "TESLA SPOEL (Basic Spark Gap Tesla Coil)", [
    "                    ┌─────────────────────────────────┐",
    "                    │         TESLA SPOEL SCHEMA       │",
    "                    └─────────────────────────────────┘",
    "",
    "  230V AC Netspanning",
    "       │",
    "  ┌────┴────┐",
    "  │  NST    │  ← Neon Sign Transformer (9kV / 30mA)",
    "  └────┬────┘",
    "       │",
    "  ┌────┴────────────────────┐",
    "  │   PRIMAIRE CIRCUIT      │",
    "  │                         │",
    "  │  C1 ══╗    ┌──────┐    │",
    "  │ (MMC) ║────┤ SPARK├────│──── Primaire spoel (L1)",
    "  │       ╚════┤  GAP ├    │     5-10 windingen 6mm Cu",
    "  │            └──────┘    │     Diameter: ~30cm",
    "  └─────────────────────────┘",
    "               │",
    "         ┌─────┴──────┐",
    "         │  SECUNDAIRE │  ← Secundaire spoel (L2)",
    "         │   SPOEL     │    800-1200 windingen 0.5mm Cu",
    "         │             │    Hoogte: 50-60cm",
    "         │   ║║║║║║║   │    Diameter: 10cm PVC buis",
    "         │   ║║║║║║║   │",
    "         │   ║║║║║║║   │",
    "         └─────┬──────┘",
    "               │",
    "          ┌────┴────┐",
    "          │ TOROID  │  ← Aluminium toroide 25x8cm",
    "          │  (top)  │    (laadt op, ontlaadt als bliksem)",
    "          └─────────┘",
    "               │",
    "           GRONDPEN (aarding verplicht!)",
    "",
    "  RESONANTIEFORMULE: f = 1 / (2π √(L·C))",
    "  Streeffrequentie: 200-400 kHz",
])

materials_box(doc, [
    "Neon Sign Transformer (NST) 9kV/30mA of 12kV/30mA",
    "MMC capacitorbank: 10x 0.1µF/2kV polypropyleencondensatoren in serie",
    "Vonkbrug: 2x wolfraamelektroden Ø6mm, verstelbare afstand 3-8mm",
    "Primaire spoel: 6mm koperen buis of kabel, 8-10 windingen, Ø30cm",
    "Secundaire spoel: PVC buis Ø10cm x 60cm, 1000 windingen 0.5mm emaildraad",
    "Toroide: aluminium luchtkanaalslang Ø8cm, gevorm tot ring Ø25cm",
    "Grondkabel: 6mm² geel/groen naar aardpen",
    "Houten of PVC basis (geen metaal!)",
    "Veiligheidsschakelaar met noodstop",
])

steps_box(doc, [
    "Wikkel de secundaire spoel: PVC buis schoonmaken, lak aanbrengen, 1000 windingen emaildraad in één richting wikkelen, afdekken met epoxy coating.",
    "Bouw de primaire spoel: 8 windingen 6mm koperen buis in spiraal, middelpunt = aansluiting vonkbrug.",
    "Maak de MMC capacitorbank: condensatoren in serie solderen, elke cap parallel met 10MΩ veiligheidsresistor.",
    "Installeer de vonkbrug: twee wolfraamelektroden op verstelbare houder, startafstand 6mm.",
    "Verbind NST → MMC → vonkbrug → primaire spoel (gesloten circuit).",
    "Verbind secundaire spoel onderkant naar aardpen, bovenkant naar toroide.",
    "Pas primaire spoel af (meer/minder windingen) voor resonantie met secundaire.",
    "EERSTE TEST: korte testpuls, controleer vonkbrug, meetapparaat op toroide.",
    "Optimaliseer: vonkbrug dichter voor hogere output, pas MMC capaciteit aan voor resonantie.",
])

body(doc, "⚠️ VEILIGHEID: Tesla-spoelen genereren dodelijke hoogspanning. Nooit alleen werken. Aarding verplicht. RF-afscherming aanbevolen. Minimale veiligheidsafstand: 3 meter bij bedrijf.")

quote(doc, "If you want to find the secrets of the universe, think in terms of energy, frequency and vibration.", "Nikola Tesla")
doc.add_page_break()

# TESLA BLUEPRINT 2 — Wardenclyffe principe
blueprint_box(doc, "WARDENCLYFFE TOWER — Resonant Earth Transmission Principe", [
    "  WARDENCLYFFE TRANSMISSIE PRINCIPE (vereenvoudigd)",
    "",
    "         ┌──────────────────┐",
    "         │  HEMISFEER/DOME  │  ← Geleidende koepel Ø10-100m",
    "         │   (TOPLOAD)      │    Laadt elektrostatisch op",
    "         └────────┬─────────┘",
    "                  │  Mast",
    "                  │  (oorspronkelijk 57m staal)",
    "         ┌────────┴─────────┐",
    "         │   GENERATOR     │  ← Hoge-freq. resonantiegenerator",
    "         │   KAMER         │    Afstemming op aardresonantie",
    "         └────────┬─────────┘",
    "                  │",
    "    ══════════════╪══════════════  Grondniveau",
    "                  │",
    "         ┌────────┴─────────┐",
    "         │  AARDSPIES       │  ← Diepte: 36m (origineel)",
    "         │  NETWERK         │    Koperen geleidende platen",
    "         │  (GRONDPLATEN)   │    Koppeling met ionosfeer via aarde",
    "         └──────────────────┘",
    "",
    "  PRINCIPE: Aarde + ionosfeer = giant resonant cavity",
    "  Schumann resonantie: ~7.83 Hz (fundamenteel)",
    "  Tesla-frequentie: afgestemd op λ/4 aardcircumferentie",
    "  Energie: niet via elektromagnetische golven maar via",
    "           geleidende laag in aardoppervlak (grondstroom)",
])

doc.add_page_break()

# ══════════════════════ 02 REMBRANDT ══════════════════════════════════════
h2(doc, "  02    Rembrandt van Rijn — Het Licht dat Verloren Ging")
h3(doc, "Wie was hij?")
body(doc, "Rembrandt Harmenszoon van Rijn (1606–1669) wordt beschouwd als de grootste schilder van de Gouden Eeuw. Zijn oeuvre omvat meer dan 300 schilderijen, 300 etsen en 2.000 tekeningen. Moderne wetenschappelijke analyse onthult dat hij materialen en technieken gebruikte die zijn tijdgenoten niet begrepen en die pas eeuwen later verklaard konden worden.")

h3(doc, "Zijn verborgen technieken")
for item in [
    "Impasto-lagen: Tientallen verf-lagen, elke laag andere compositie — complete schilderijen onder de oppervlakte.",
    "Loodwit-experimenten: Unieke mengsels met smaliet (kobaltblauw glas) — nooit volledig gereconstrueerd.",
    "Bindmiddelmengsels: Unieke olie-hars combinaties voor ongeëvenaarde transparantie.",
    "Perceptueel licht: Schilderde wat het brein ziet, niet wat het oog ziet — 20e-eeuwse optische psychologie.",
]:
    doc.add_paragraph(f"  • {item}", style='List Bullet')

h3(doc, "Hoe ging de kennis verloren?")
body(doc, "Het gildestelsel controleerde kunstenaars via commissies en strikte materiaalregels. Rembrandt weigerde zich aan te passen en ging failliet in 1656. Zijn formules leefden in zijn handen — niet in documenten.")

divider(doc)

blueprint_box(doc, "REMBRANDT IMPASTO TECHNIEK — Laagopbouw Schema", [
    "  LAAGOPBOUW EEN REMBRANDT PORTRET (cross-section)",
    "",
    "  Laag 7: Vernislaag               ░░░░░░░░░░░░░░░░░░░░░░░  (0.05mm)",
    "  Laag 6: Glazuurlaag (transparant) ▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒  (0.1mm)",
    "  Laag 5: Impasto hoogtepunten     ████████████████████████  (1-3mm!)",
    "  Laag 4: Modelleerlaag vlees      ▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓▓  (0.3mm)",
    "  Laag 3: Kleurlaag                ████████████████████████  (0.2mm)",
    "  Laag 2: Grijze imprimatura       ▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒  (0.1mm)",
    "  Laag 1: Grond (loodwit+krijt)    ░░░░░░░░░░░░░░░░░░░░░░░  (0.5mm)",
    "  Drager:  Linnen of eiken paneel  ========================",
    "",
    "  REMBRANDTS PIGMENTFORMULE (gereconstrueerd door RRP):",
    "  Huidtint:  Loodwit 60% + Vermiljoen 15% + Oker 20% + Smaliet 5%",
    "  Bindmiddel: Lijnolievarnish + dammarshars 80:20",
    "  Impasto: dezelfde mix maar 3x meer loodwit, dikker gemalen",
    "",
    "  LICHTPRINCIPE (chiaroscuro):",
    "  Donker vlak ──────► midtoon ──────► highlight (impasto)",
    "  Dunne lagen            medium         Dikke paste",
    "  (transparant)                        (reflecteert direct licht)",
])

materials_box(doc, [
    "Linnen canvas of eiken paneel (geprepareerd met krijt-lijnolieprimer)",
    "Loodwit (gebruik zinkwit als vervanger vanwege toxiciteit loodwit)",
    "Vermiljoen (cadmiumrood als moderne vervanger)",
    "Geelochre en bruinocher pigmenten",
    "Smaliet (kobaltblauw glas, fijngemalen) — optioneel, voor koude tonen",
    "Lijnolievarnish (gezuiverd, zongebleekt)",
    "Dammarshars (opgelost in terpentijn 1:3)",
    "Verf-spatels voor impasto-applicatie",
    "Zachte borstels voor glazuurlagen",
])

steps_box(doc, [
    "Bereid de drager voor: lijn canvas op spieraam, breng 2-3 lagen krijt-lijnolieprimer aan.",
    "Teken de compositie in houtskool, fix met fixeerspray.",
    "Breng de grijze imprimatura aan: dunne laag loodwit gemengd met umber (warme toon) of grijs.",
    "Onderschilder dood (dead layer): volledig uitwerken in grijstonen, alle vormen definiëren.",
    "Eerste kleurlaag: dunne verf, transparant, warm in de schaduwen.",
    "Modelleerlaag: opbouwen van de vormen met medium-dikke verf, warm vlees voor huid.",
    "Impasto highlights: de lichtste lichten met dikke pasta direct uit tube/spatel, niet uitvegen.",
    "Laat 6-12 maanden drogen (traditioneel — moderne alkydmedium versnelt dit naar 1 week).",
    "Glazuurlagen: ultra-dunne transparante lagen voor diepte en glans in schaduwen.",
    "Vernis: na volledig drogen (6-12 maanden) met dunne dammarvernislaag.",
])

doc.add_page_break()

# ══════════════════════ 03 RIFE ════════════════════════════════════════════
h2(doc, "  03    Royal Raymond Rife — Frequenties als Medicijn")
h3(doc, "Wie was hij?")
body(doc, "Royal Raymond Rife (1888–1971) was een Amerikaanse microbiologisch onderzoeker die in de jaren 1920-1930 een universele microscoop ontwikkelde met een resolutie die decennia later pas door elektronenmicroscopen geëvenaard werd. Zijn bewering: levende virussen en bacteriën konden worden vernietigd met specifieke elektromagnetische resonantiefrequenties — de Mortal Oscillatory Rate (MOR). In 1934 rapporteerden alle 16 terminale kankerpatiënten in een klinische studie volledige genezing.")

h3(doc, "De vernietigingscampagne")
body(doc, "Morris Fishbein van de American Medical Association startte een systematische campagne. Samenwerkende artsen verloren hun licenties. Rife's laboratorium werd vernield. Zijn financiers werden uitgeschakeld. Rife stierf arm en verlaten.")

divider(doc)

blueprint_box(doc, "RIFE FREQUENTIEGENERATOR — Basis Solid-State Circuit", [
    "  RIFE-GEBASEERDE FREQUENTIEGENERATOR (moderne replica)",
    "",
    "  ┌─────────────┐    ┌────────────┐    ┌──────────────┐",
    "  │  OSCILLATOR │    │  VERSTERKER│    │  OUTPUT      │",
    "  │  (DDS/555)  │───►│  (MOSFET)  │───►│  ELEKTRODEN  │",
    "  └─────────────┘    └────────────┘    └──────────────┘",
    "         │                  │                  │",
    "  ┌─────────────┐    ┌────────────┐    ┌──────────────┐",
    "  │  Arduino/   │    │  IRF530 of │    │  RVS of Koper│",
    "  │  Raspberry  │    │  IRF540N   │    │  elektroden  │",
    "  │  Pi Zero    │    │  MOSFET    │    │  (20x5cm)    │",
    "  └─────────────┘    └────────────┘    └──────────────┘",
    "",
    "  RIFE MOR FREQUENTIES (historisch gedocumenteerd):",
    "  Bacillus X (kanker): 11,780,000 Hz (11.78 MHz)",
    "  BX virus:             2,128 Hz  (audiofrequentie)",
    "  Staphylococcus:        727 Hz",
    "  Streptococcus:         880 Hz",
    "  E. coli:               802 Hz",
    "  Candida:              2,008 Hz",
    "",
    "  MODERNE EQUIVALENTEN (hedendaags onderzoek):",
    "  PEMF apparaten: 1-100 Hz pulserend veld",
    "  HIFU (High Intensity Focused Ultrasound): 1-3 MHz",
    "  Cymatic therapie: 30-300 Hz",
    "",
    "  CIRCUIT (Arduino-gebaseerd):",
    "  Arduino D9 (PWM) ──────► Gate IRF540N ──────► Elektrode A",
    "  Arduino GND ────────────► Source IRF540N ─────► Gemeenschappelijk",
    "  5V ─── 10kΩ resistor ──► Gate (pull-up)",
    "  Aanpasbare frequentie: 1 Hz tot 100 kHz via software",
])

materials_box(doc, [
    "Arduino Uno of Nano (frequentiegenerator via PWM of DDS-module AD9833)",
    "DDS-module AD9833 (voor precisiefrequenties tot 12.5 MHz)",
    "IRF540N N-channel MOSFET (of IRF530 voor lagere vermogens)",
    "10kΩ weerstand (gate pull-down)",
    "Voeding 12V/2A (voor elektrodenstroom)",
    "Roestvrij stalen elektroden 20x5cm (of koperen plaatjes)",
    "Geleidend gel of zoutoplossing (voor huid-elektrodecontact)",
    "Oscilloscoop of freq-meter voor verificatie",
    "LCD-display voor frequentieweergave (optioneel)",
])

steps_box(doc, [
    "Download Arduino IDE. Installeer DDS-library (MD_AD9833 of Adafruit)",
    "Laad de frequentiegenerator sketch: stel frequentie in via seriële monitor.",
    "Verbind AD9833 module: VCC→5V, GND→GND, DATA/CLK/FSYNC naar SPI-pinnen.",
    "Verbind MOSFET: Gate→D9(PWM) of DDS output, Drain→elektrode+, Source→GND.",
    "Verbind 12V voeding Drain-kant via 100Ω stroombeperker.",
    "Gebruik oscilloscoop om output-golf te verifiëren (sinusvorm of blokgolf).",
    "Test met zoutoplossing-bakje eerst (geen directe huidcontact) — meet stroomdichtheid.",
    "Raadpleeg voor therapeutisch gebruik altijd een arts. Dit is uitsluitend voor onderzoek.",
])

body(doc, "⚠️ DISCLAIMER: Dit circuit is voor educatief en onderzoeksdoeleinden. Geen medische claims. Raadpleeg altijd een arts voor gezondheidsklachten.")
doc.add_page_break()

# ══════════════════════ 04 SEMMELWEIS ═════════════════════════════════════
h2(doc, "  04    Ignaz Semmelweis — Handen Wassen Redde Levens, maar Kostte Hem Alles")
body(doc, "Ignaz Semmelweis (1818–1865) was een Hongaarse gynaecoloog werkzaam in Wenen. In 1847 ontdekte hij dat kraamvrouwenkoorts sterfte daalde van 18% naar onder 2% wanneer artsen handen wasten met chlooroplossing vóór bevallingen. De implicatie — artsen doodden hun patiënten — was voor de medische wereld onaanvaardbaar. Hij werd in een psychiatrische instelling opgesloten en stierf aan een infectie.")

divider(doc)

blueprint_box(doc, "SEMMELWEIS PROTOCOL — Gereconstrueerd Aseptisch Handwasprotocol", [
    "  SEMMELWEIS CHLOORWAS PROTOCOL (1847)",
    "",
    "  INGREDIËNTEN:",
    "  ┌────────────────────────────────────────────────┐",
    "  │  Calciumhypochloriet (bleekpoeder/chloorkalk)  │",
    "  │  Chloride of lime: Ca(ClO)₂                   │",
    "  │  Concentratie oplossing: ~1-2% actief chloor  │",
    "  │  Bereiding: 10g bleekpoeder per liter water   │",
    "  └────────────────────────────────────────────────┘",
    "",
    "  PROTOCOL (Semmelweis 1847):",
    "  1. Na elke lijkschouwing of sectie:",
    "     → Handen weken in chlooroplossing: 2-3 minuten",
    "     → Alle vingers, nagels, handpalmen",
    "     → Daarna afspoelen met schoon water",
    "  2. Vóór elke geboorte/gynaecologisch onderzoek:",
    "     → Herhaal wasstap",
    "     → Gebruik schone handdoek (eenmalig gebruik)",
    "",
    "  MODERNE VERTALING (CDC hand-hygiëne protocol):",
    "  ┌────────────────────────────────────────────────┐",
    "  │ Alcoholische handgel (70% ethanol/isopropanol) │",
    "  │ OF 20 sec wassen met zeep + water              │",
    "  │ Mechanische wrijving: alle oppervlakken 20 sec │",
    "  └────────────────────────────────────────────────┘",
    "",
    "  WERKINGSMECHANISME:",
    "  HOCl (hypochloorzuur) → oxidatie bacteriecelmembraan",
    "  → eiwitdenaturatie → bacteriedood in <30 seconden",
    "  Effectief tegen: Gram+ en Gram- bacteriën, virussen",
])

doc.add_page_break()

# ══════════════════════ 05 REICH ══════════════════════════════════════════
h2(doc, "  05    Wilhelm Reich — Het Lichaam als Archief van Trauma")
body(doc, "Wilhelm Reich (1897–1957) was een Oostenrijks-Amerikaanse psychoanalyticus. Hij ontwikkelde de theorie van karakterpantsering — lichamelijke spierspanning als fysieke neerslag van psychologische trauma's. In 1956 verbrandde de Amerikaanse FDA letterlijk zijn boeken. Reich stierf in de federale gevangenis van Lewisburg in 1957.")

divider(doc)

blueprint_box(doc, "ORGON ACCUMULATOR — Reich's Energie-Condenser", [
    "  ORGON ACCUMULATOR — CONSTRUCTIESCHEMA",
    "",
    "  PRINCIPE: alternerende organische (isolerende) en metalen",
    "  (geleidende) lagen creëren een 'condenser' voor orgonenergie.",
    "  Reich beweerde dat meer lagen = sterkere accumulatie.",
    "",
    "  DOOS-ACCUMULATOR (5-laags, voor persoonlijk gebruik):",
    "",
    "  Buitenste laag → Hout (organisch)           20mm dik",
    "               → Staalwol/metaalfolie         2mm dik",
    "               → Multiplex (organisch)         10mm dik",
    "               → Aluminium folie               1mm dik",
    "               → Binnenste hout (organisch)    10mm dik",
    "               → BINNENRUIMTE ← persoon zit hier",
    "",
    "  AFMETINGEN BOUW-INSTRUCTIE:",
    "  Buiten: 100cm breed × 100cm diep × 180cm hoog (staand)",
    "  Deur: scharnierend 90x180cm",
    "  Stoel: houten stoel, geen metaal",
    "",
    "  MATERIAALSTAPEL (per wand, van buiten naar binnen):",
    "  [Hout 20mm] [Staalwol 3mm] [Hout 10mm] [Aluminiumfolie] [Hout 10mm]",
    "",
    "  ORGONITE (moderne variant):",
    "  Epoxyhars 50% + Metaalvijlsel 50% + Kwartskristal",
    "  → Gegoten in mallen, uitgehard 24-48 uur",
    "  → Moderne 'orgoniet' pyramides volgen hetzelfde principe",
])

materials_box(doc, [
    "Multiplex platen 18mm (buiten- en binnenlagen)",
    "Staalwol fijn (bos) — GEEN gegalvaniseerd of roestvrij staal",
    "Aluminium folie (keukenfolie volstaat) of dun aluminiumplaat",
    "Houtlijm en schroeven",
    "Deurscharnieren (kunststof, geen metaal)",
    "Houten stoel voor binnenin",
    "Voor orgoniet: vloeibare epoxyhars + stalen/koper vijlsel + kwarts",
])

steps_box(doc, [
    "Zaag multiplexplaten op maat voor bodem, zijwanden, plafond, achterwand.",
    "Leg staalwol in dunne, gelijkmatige laag op elk multiplexpanel.",
    "Bedek staalwollaag met aluminium folie, goed aandrukken.",
    "Bevestig volgende houtlaag bovenop (sandwich). Herhaal 2-3x per wand.",
    "Assembleer de doos met schroeven en houtlijm. Deur met kunststofscharnier.",
    "Binnenin: alleen houten stoel, geen metalen objecten.",
    "Gebruik: 30-60 minuten per sessie, niet langer (Reich's eigen richtlijn).",
    "Ontlading: ventileer dagelijks, bewaar niet in de zon (oververhitting).",
])

doc.add_page_break()

# ══════════════════════ 06 HILDEGARD ══════════════════════════════════════
h2(doc, "  06    Hildegard van Bingen — De Polymath die 700 Jaar Werd Vergeten")
body(doc, "Hildegard van Bingen (1098–1179) was een Duitse abdis, componiste, schrijfster, mystica en genezer. Haar oeuvre omvat theologie, natuurwetenschap, geneeskunde, muziek en kosmologie. Ze schreef over kruidengeneeskunde en psychosomatische verbanden in een tijd dat vrouwen geacht werden te bidden en te zwijgen. Ze werd heilig verklaard in 2012.")

divider(doc)

blueprint_box(doc, "HILDEGARD'S KRUIDENTHERAPIE — Gereconstrueerde Formules", [
    "  HILDEGARD VAN BINGEN — PHYSICA RECEPTEN (Causae et Curae)",
    "",
    "  1. SPELTBROOD RECEPT (digestie & energie):",
    "     Ingrediënten: Spelt 500g, water 300ml, zout 8g, gist 7g",
    "     Methode: kneden 10 min → rijzen 60 min → bakken 220°C/35 min",
    "     Hildegard: 'Spelta is de beste graansoort, warmt het bloed'",
    "     Modern: hoog magnesium, B-vitamines, FODMAP-vriendelijker dan tarwe",
    "",
    "  2. LAVENDEL WIJN (hoofdpijn, melancholie):",
    "     100ml lavendelbloesem + 1L witte wijn + 2 eetl honing",
    "     7 dagen macereren, filteren",
    "     Dosis: 1 kleine glas voor het slapen",
    "     Modern: lavendel = linalool (anxiolyticum, bewezen in studies)",
    "",
    "  3. HILDEGARD KOEKJES (kerstkoekjes — 'Nerve cookies'):",
    "     Speculaaskruiden 10g + Muskaatbloem 3g + Kaneel 5g",
    "     Nootmuskaat 2g + Nelkenkruid 1g + Speltmeel 300g",
    "     Honing 100g + Boter 150g",
    "     → Gemengd, uitgerold 5mm, gebakken 180°C/12 minuten",
    "     Hildegard: 'verlicht de sombere geest en verblijdt het hart'",
    "",
    "  4. FENEGRIEK TINCTURE (leverondersteuning):",
    "     50g fenegriekzaden + 500ml appelazijn",
    "     14 dagen macereren, dagelijks schudden",
    "     Filteren, bewaren in donkere fles",
    "     Dosis: 1 theelepel voor de maaltijd",
    "     Modern: fenegriek = trigonelline, saponines — glycemische controle",
    "",
    "  VIRIDITAS — HET LEVENSBEGINSEL:",
    "  Hildegard's concept: planten/mensen bevatten 'groenkracht'",
    "  Modern parallel: chlorofyl + mitochondriële energieproductie",
    "  Praktijk: dagelijks vers groen (rucolasla, peterselie, basilicum)",
])

doc.add_page_break()

# ══════════════════════ NIEUWE UITVINDERS ══════════════════════════════════
h1(doc, "DEEL II — AANVULLENDE GENIALE UITVINDERS VAN VERBORGEN KENNIS")
doc.add_page_break()

# 07 — Viktor Schauberger
h2(doc, "  07    Viktor Schauberger — De Taal van Levend Water")
h3(doc, "Wie was hij?")
body(doc, "Viktor Schauberger (1885–1958) was een Oostenrijkse boswachter en uitvinder die zijn leven wijdde aan het begrijpen van water als levend medium. Zonder formele wetenschappelijke opleiding ontwikkelde hij een complete theorie van imploderende (naar binnen gerichte) beweging als basis voor vrije energie en antigravitatie — het tegendeel van de explosive beweging waarop al onze technologie gebaseerd is.")
body(doc, "Zijn uitvindingen omvatten zelfreinigende waterbuizen die water verkoelen in plaats van verwarmen, vortex-turbines die energie winnen uit waterbewegingen, en beweerde anti-zwaartekrachtapparaten gebaseerd op implosie. De nazi's dwongen hem vliegende-schijf-prototypen te bouwen. Na de oorlog lokten Amerikanen hem naar de VS — hij tekende patenten over, en stierf 5 dagen na thuiskomst in Oostenrijk.")

h3(doc, "Zijn ontdekkingen")
for item in [
    "Implosie vs Explosie: Alle natuur beweegt in spiralen naar binnen (implosie). Onze technologie gaat explosief naar buiten — het tegenovergestelde van levende systemen.",
    "Vortex waterbehandeling: Water in een spiraalvorm bewegen langs de as van een buis restaureert zijn 'levende' eigenschappen.",
    "Logaritmische spiraal (Fibonacci): De structuur van waterbewegingen volgt exact de Fibonacci-spiraal.",
    "Trout Turbine: Een turbine die de spiraalbeweging van een forel nabootst om energie te winnen.",
    "PKS Pijp: Koperen pijpen met spiraalgroeven die water koel, snel en levend houden — nu erkend in hydrologisch onderzoek.",
    "Repulsin: Schijfvormig apparaat dat lucht in een spiraalpatroon beweegt voor heffing — de nazi's testten dit als vliegende schijf.",
]:
    doc.add_paragraph(f"  • {item}", style='List Bullet')

h3(doc, "Hoe werd het onderdrukt?")
body(doc, "De SS dwong Schauberger in 1943 zijn Repulsin te bouwen. Na de oorlog nam de Amerikaanse inlichtingendienst zijn werk over. Hij werd naar de VS gelokt door Texan investors, gedwongen zijn kennis over te dragen, en stierf kort daarna onder verdachte omstandigheden. Al zijn papieren zijn in handen van de Oostenrijkse 'Schauberger Foundation', maar de technische details van zijn Repulsin zijn nooit vrijgegeven.")

divider(doc)

blueprint_box(doc, "SCHAUBERGER VORTEX PIJP — Imploderende Waterbeweging", [
    "  SCHAUBERGER PKS VORTEX PIJP — SPIRAALSTROOM PRINCIPE",
    "",
    "  DOORSNEDE — SPIRAALSTROOM IN BUIS:",
    "",
    "  Conventionele buis:          Schauberger vortex buis:",
    "  Rechtlijnig stromen           Spiraalvormig stromen",
    "     → → → → →                    ↙ ↗ ↙ ↗ ↙",
    "  Wrijving: HOOG                Wrijving: LAAG",
    "  Bacteriegroei: Ja             Bacteriegroei: Nee",
    "  Temperatuur: stijgt           Temperatuur: daalt",
    "",
    "  SCHAUBERGER EI-VLOEISTOF OPSLAG:",
    "  ┌──────────────────────────────────┐",
    "  │          ___________              │",
    "  │         /           \\             │",
    "  │        /   EI-VORM   \\            │",
    "  │       │   (koperen   │            │",
    "  │       │    vat)      │            │",
    "  │        \\             /            │",
    "  │         \\___________/             │",
    "  │               │                  │",
    "  │           smalle punt            │",
    "  │         (vortexpunt onder)       │",
    "  └──────────────────────────────────┘",
    "  Schauberger's bronwater-opslag:",
    "  Ei-vorm → continu interne vortexbeweging → water blijft 'levend'",
    "",
    "  VORTEX RATIO:",
    "  Optimale spiraalhoek: 62° (gouden hoek, ≈ Fibonacci spiraal)",
    "  Aantal spiraalwendingen per diameter: 1.618 (gouden ratio φ)",
])

materials_box(doc, [
    "Koperen buis Ø22mm × 2m (koper is Schauberger's voorkeursmateriaal)",
    "Spiraalvorm mal: hout of 3D-print, diameter variabel",
    "Soldeertin loodvrij (voor verbindingen)",
    "Heatgun voor het buigen van koper",
    "Ei-vormige keramische pot (klei, geen plastic) voor wateropslag",
    "Digitale thermometer voor voor/na temperatuurmeting",
])

steps_box(doc, [
    "Verhit koperen buis gelijkmatig met heatgun tot ductiel (zachter).",
    "Buig de buis in een spiraalvorm rond de mal: 62° hellingshoek per winding.",
    "Eindig in een recht stuk voor in- en uitlaat.",
    "Laat afkoelen, verwijder mal. Controleer op knikken.",
    "Verbind aan watertoevoer. Meet inkomende temperatuur.",
    "Meet na 5 minuten stromen de uitkomende temperatuur.",
    "Verwacht: water koelt 0.5-2°C — documenteer dit als controle-experiment.",
    "Vergelijk met rechte buis van zelfde lengte (controletest).",
])

quote(doc, "Nature is not served by rigid laws, but by rhythmic, life-giving processes.", "Viktor Schauberger")
doc.add_page_break()

# 08 — Walter Russell
h2(doc, "  08    Walter Russell — De Kosmos als Levende Geometrie")
h3(doc, "Wie was hij?")
body(doc, "Walter Russell (1871–1963) was een Amerikaans wetenschapper, kunstenaar, architect, filosoof en muzikant — een van de meest universele geesten van de 20e eeuw. Zonder formele wetenschappelijke opleiding presenteerde hij in 1926 een compleet alternatief voor het periodiek systeem van elementen, inclusief voorspellingen van elementen die pas decennia later ontdekt werden (deuterium, plutonium, tritium). Hij beweerde zijn inzichten te hebben ontvangen in een 39-daagse mystieke verlichtingservaring in 1921.")

h3(doc, "Zijn ontdekkingen")
for item in [
    "Uitgebreid Periodiek Systeem: Voorspelde in 1926 bestaan van deuterium (bevestigd 1931), plutonium en andere transuranische elementen.",
    "Transmutatie van elementen: Beweerde dat elementen kunnen transformeren onder de juiste elektromagnetische condities — destijds onmogelijk geacht, nu nucleaire basis.",
    "Ruimte als medium: Stelde dat ruimte niet leeg is maar een drukmedium — vergelijkbaar met het moderne begrip van het kwantumvacuüm.",
    "Octave Wave Principle: Alle materie beweegt in octaven van golffrequenties — parallellen met moderne snaartheorie.",
    "Vortex-gebaseerde kosmologie: Sterrenstelsels en atomen volgen identieke vortexpatronen — fractale kosmologie.",
]:
    doc.add_paragraph(f"  • {item}", style='List Bullet')

h3(doc, "Hoe werd het onderdrukt?")
body(doc, "Russell's werk werd door de wetenschappelijke gemeenschap genegeerd omdat het buiten alle gevestigde kaders viel. Hij had geen peer-reviewed publicaties in erkende tijdschriften. Ironisch: zijn voorspellingen van ontbrekende elementen kwamen uit — zijn methode werd nooit serieus onderzocht. Zijn werk overleeft via de 'University of Science and Philosophy' die hij in 1948 met zijn vrouw Lao Russell oprichtte in Swannanoa, Virginia.")

divider(doc)

blueprint_box(doc, "RUSSELL'S PERIODIEK SYSTEEM UITBREIDING — Octave Wave Schema", [
    "  RUSSELL'S OCTAVE WAVE PERIODIC TABLE (1926 vereenvoudigd)",
    "",
    "  Standaard periodiek systeem: 7 perioden, 118 elementen",
    "  Russell's uitbreiding: 9 octaven, geometrische plaatsing",
    "",
    "  RUSSELL'S ELEMENT-OCTAVEN:",
    "  Octaaf 1: H  He",
    "  Octaaf 2: Li Be B  C  N  O  F  Ne",
    "  Octaaf 3: Na Mg Al Si P  S  Cl Ar",
    "  Octaaf 4: K  Ca Sc Ti V  Cr Mn Fe Co Ni Cu Zn Ga Ge As Se Br Kr",
    "  Octaaf 5: Rb Sr ...",
    "  Octaaf 9: [onontdekte elementen > 118] ← Russell voorspelde deze",
    "",
    "  RUSSELL'S KOSMISCHE GEOMETRIE:",
    "",
    "       YANG (positief, samen-trekkend)  YING (negatief, uitzettend)",
    "             \\                    /",
    "              \\                  /",
    "               \\     GOLF      /",
    "                \\    KNOOP   /",
    "                 \\   (materie)/",
    "                  \\        /",
    "                   X──────X  ← ELEMENT POSITIE",
    "                  /        \\",
    "                 /  LEEGTE  \\",
    "                /   (ruimte) \\",
    "",
    "  Russell: 'Materie is niet — het wordt voortdurend'",
    "  Modern parallel: kwantumveld — deeltjes als excitaties van het veld",
])

doc.add_page_break()

# 09 — Thomas Townsend Brown
h2(doc, "  09    Thomas Townsend Brown — Zwaartekracht Controleren met Electriciteit")
h3(doc, "Wie was hij?")
body(doc, "Thomas Townsend Brown (1905–1985) was een Amerikaanse wetenschapper en marinofficier die in 1921, als 16-jarige, ontdekte dat een sterk geladen condensator neiging heeft te bewegen in de richting van zijn positieve pool — het Biefeld-Brown effect. Zijn leven lang claimde hij dat dit niet alleen elektrostatisch was maar een koppeling tussen elektriciteit en zwaartekracht aantoonde — een fundamentele herziening van de natuurkunde.")

h3(doc, "Zijn ontdekkingen")
for item in [
    "Biefeld-Brown Effect: Geladen asymmetrische condensatoren bewegen naar hun positieve pool — in vacuum bewezen (niet alleen 'ionenwind').",
    "Gravitoelectrics: Theorie dat zware massa's een zwak elektrisch veld genereren en dat dit omgekeerd ook werkt.",
    "Electrokinetics: Verzamelnaam voor zijn pogingen om elektrische aandrijving zonder bewegende delen te bouwen.",
    "Onderzeese detectie: Brown's geostationaire gravito-elektrische detector werd gebruikt door de US Navy voor onderzeeër-detectie — het enige officieel erkende gebruik van zijn werk.",
    "Saucer-platforms: In 1955 demonstreerde Brown aan de VS-luchtmacht vliegende 90cm-schijven aangedreven door hoogspanning (50kV) — zonder brandstof.",
]:
    doc.add_paragraph(f"  • {item}", style='List Bullet')

h3(doc, "Hoe werd het onderdrukt?")
body(doc, "Na zijn demonstratie in 1955 werd Brown's werk geclassificeerd door de US Air Force. Project Winterhaven — zijn voorstel voor een volledig elektrisch aangedreven vliegtuig — werd afgewezen. De officiële reden: 'ionenwind verklaring volstaat.' Veel onderzoekers stellen dat vacuum-tests die ionenwind uitsluiten wel degelijk een restkracht tonen. De documentatie van Project Winterhaven is nooit volledig vrijgegeven.")

divider(doc)

blueprint_box(doc, "BIEFELD-BROWN LIFTER — Elektrostatische Heffing", [
    "  BIEFELD-BROWN LIFTER — ASYMMETRISCHE CONDENSATOR",
    "",
    "  BOVENAANZICHT:",
    "         Dunne draad (positief, +30kV)",
    "              ┌─────────────────────┐",
    "              │  ─────────────────  │  ← 0.1mm koperdraad",
    "              │                     │",
    "              │     LUCHTGAP        │  ← 3-5cm scheiding",
    "              │     (diëlektricum)  │",
    "              │                     │",
    "              │  ─────────────────  │  ← aluminium folie (negatief, GND)",
    "              └─────────────────────┘",
    "        Aluminium folierand (negatief, GND/aarde)",
    "",
    "  ZIJAANZICHT:",
    "  +30kV ──── dunne draad ────────────┐",
    "                                      │ luchtspatie 3-5cm",
    "  GND ──── aluminiumfolie plaat ─────┘",
    "                │",
    "                ▼ (bewegingsrichting: draad=positief pool)",
    "",
    "  SPECIFICATIES VOOR DEMO-LIFTER:",
    "  Draad: 0.1mm koperdraad of gitaarsnaar (dun)",
    "  Folie: aluminiumfolie 20x5cm",
    "  Scheiding: 4cm (optimaal)",
    "  Spanning: 25-40kV DC",
    "  Stroom: < 1mA (micro-amp niveau)",
    "  Gewicht van lifter: < 2 gram voor heffing bij 30kV",
    "",
    "  KRACHTFORMULE (Biefeld-Brown, empirisch):",
    "  F = k × V² × A / d²",
    "  F=kracht, V=spanning, A=condensatoroppervlak, d=afstand",
])

materials_box(doc, [
    "Spanningsbron 25-40kV DC (Neon Sign Transformer + gelijkrichter, of vliegendverdelger HV-module)",
    "0.1mm koperdraad (ca. 30cm per lifter-ring)",
    "Aluminium folie (standaard keukenfolie)",
    "Balsahout of Depron foam (structuur, ultralicht)",
    "Hoge-spanning silicone draden (10kV isolatie)",
    "HV-klemmverbindingen",
    "Veiligheidsafscherming (plexi-glas ombouw verplicht!)",
])

steps_box(doc, [
    "Bouw driehoekige of rechthoekige frame van balsahout, ca. 20x20cm.",
    "Span dunne koperdraad langs bovenste rand van het frame.",
    "Bevestig aluminiumfolie langs onderste rand (3-5cm onder draad).",
    "Verbind draad aan +HV (30kV), folierand aan GND via HV-kabels.",
    "Hang lifter op met dun katoen draad (niet geleidend) of laat vrij staan op insulator.",
    "Schakel HV in van afstand (minimaal 1 meter veiligheidsafstand).",
    "Observeer heffing — bij correct gebouwde lifter zichtbaar bewegen.",
    "Test ook in afgesloten doos (verminderde ionenwind) voor wetenschappelijke controle.",
])

body(doc, "⚠️ VEILIGHEID: 30kV is DODELIJK. Gebruik altijd isolation transformer. NOOIT aanraken bij bedrijf. Bouw enkel als je ervaring hebt met hoogspanningswerk.")

quote(doc, "The gravitational field is fundamentally an electromagnetic phenomenon.", "T.T. Brown")
doc.add_page_break()

# 10 — Georges Lakhovsky
h2(doc, "  10    Georges Lakhovsky — De Frequentie van Leven")
h3(doc, "Wie was hij?")
body(doc, "Georges Lakhovsky (1869–1942) was een Russisch-Franse ingenieur en biofysicus die in de jaren 1920-1930 een revolutionaire theorie van leven ontwikkelde: elke cel in het lichaam is een oscillerende circuit — een kleine antenne die uitzendt en ontvangt op zijn eigen resonantiefrequentie. Ziekte is verstoring van die resonantie; gezondheid is harmonische oscillatie.")
body(doc, "Zijn uitvinding, de Multiple Wave Oscillator (MWO), zond een volledig spectrum van frequenties uit — en stimuleerde cellen om op hun eigen resonantiefrequentie terug te keren. Rapporten uit de jaren 1930 van het Hôpital de la Salpêtrière in Parijs beschreven succesvolle behandeling van planten, dieren en mensen.")

h3(doc, "Zijn ontdekkingen")
for item in [
    "Cellulaire Oscillatie: Elke cel is een LC-circuit (spoel + condensator) dat oscilleert op zijn eigen frequentie.",
    "Multiple Wave Oscillator: Apparaat dat gelijktijdig duizenden frequenties uitzendt — spectrum 750kHz tot 3GHz.",
    "Cosmic Rays & Life: Theorie dat kosmische straling de primaire energiebron is voor biologische oscillatie.",
    "Boomplant experimenten: Planten met kankerachtige tumoren genezen door MWO-behandeling — gedocumenteerd in foto's (1925-1931).",
    "The Secret of Life (1925): Zijn publicatie die de cel-oscillatietheorie beschrijft, gepubliceerd in meerdere Europese talen.",
]:
    doc.add_paragraph(f"  • {item}", style='List Bullet')

h3(doc, "Hoe werd het onderdrukt?")
body(doc, "Lakhovsky vluchtte in 1940 naar de VS na de nazi-invasie van Frankrijk. In New York behandelde hij enkele patiënten met groot succes. In 1942 werd hij op mysterieuze wijze aangereden door een auto en overleed. Zijn apparaten en documentatie verdwenen grotendeels. Zijn boek 'The Secret of Life' is nog steeds verkrijgbaar maar zijn technische specificaties zijn verspreid en nooit officieel gerepliceerd.")

divider(doc)

blueprint_box(doc, "MULTIPLE WAVE OSCILLATOR (MWO) — Lakhovsky's Frequentie-Spectrum Antenne", [
    "  LAKHOVSKY MULTIPLE WAVE OSCILLATOR — SCHEMA",
    "",
    "  BOVENAANZICHT (één antennering-set):",
    "",
    "         ○ ○ ○ ○ ○ ○ ○ ○ ○  ← 9 concentrische ringen",
    "        ○               ○",
    "       ○                 ○   Elke ring = andere frequentie",
    "      ○                   ○  Ring 1 (buiten): laagste freq.",
    "     ○                     ○ Ring 9 (binnen): hoogste freq.",
    "      ○                   ○",
    "       ○                 ○",
    "        ○               ○",
    "         ○ ○ ○ ○ ○ ○ ○ ○ ○",
    "",
    "  ZIJAANZICHT (twee antenna-sets, patiënt ertussen):",
    "  ┌─────────────────────────────────────────────┐",
    "  │  [ZENDANTENNE]  ← PATIËNT →  [ONTVANGST]   │",
    "  │  (9 ringen)                  (9 ringen)     │",
    "  └─────────────────────────────────────────────┘",
    "                      │",
    "              Tesla-spoel of",
    "              vonkgap oscillator",
    "",
    "  RINGFORMULE (Lakhovsky):",
    "  Frequentie ring n: f_n = c / (2π × r_n × √LC)",
    "  Lakhovsky gebruikte geometrische progressie:",
    "  r_1 = 10cm, r_2 = 12cm, r_3 = 14.4cm, ...",
    "  (ratio: ×1.2 per ring)",
    "",
    "  MATERIALEN RINGEN: koperdraad Ø3mm",
    "  Open eindjes! (Lakhovsky: gesloten ringen werken niet)",
    "  Open eindjes = staan 2-3mm van elkaar = vonk mogelijk",
])

materials_box(doc, [
    "Koperdraad Ø3mm (ca. 5 meter per ring-set)",
    "Neon Sign Transformer 9-15kV als HV-bron",
    "Houten of Bakeliet frame (geen metaal voor de structuur)",
    "Epoxy isolators voor ring-bevestiging",
    "HV-kabel 10kV isolatie voor verbinding NST→ringen",
    "Vonkgap of Tesla-spoel als modulatie (optioneel)",
])

steps_box(doc, [
    "Buig 9 concentrische koperen ringen: diameters 20, 24, 29, 35, 42, 50, 60, 72, 86cm.",
    "Laat elke ring open: eindjes 5mm van elkaar, niet verbonden.",
    "Bevestig ringen met epoxy op isolerende steunen, concentrisch uitgelijnd.",
    "Bouw twee identieke ring-sets (zend- en ontvangstantenne).",
    "Verbind buitenste ring van elke set aan NST-uitgang.",
    "Stel op ca. 80cm uit elkaar (patiënt of proefobject ertussen).",
    "Test eerst op planten (niet op mensen) met 1-2 minuten per sessie.",
    "Documenteer resultaten — groei, kleur, wortelsysteem vóór en na.",
])

doc.add_page_break()

# 11 — John Worrell Keely
h2(doc, "  11    John Worrell Keely — Harmonische Resonantie als Vrije Energie")
h3(doc, "Wie was hij?")
body(doc, "John Ernst Worrell Keely (1827–1898) was een Amerikaanse uitvinder uit Philadelphia die beweerde een motor te hebben gebouwd die werkte op 'etherkracht' — aangedreven door de resonantie van muziekakkoorden. Zijn Keely Motor Company trok miljoenen dollars aan investeringen, maar produceerde nooit een werkende commerciële motor. Was hij een geniaal experimentator of de meest briljante oplichter van de 19e eeuw? Het antwoord is — waarschijnlijk beide.")

h3(doc, "Zijn beweerde ontdekkingen")
for item in [
    "Sympathetic Vibratory Physics: Theorie dat alle materie trilt en dat materie gemanipuleerd kan worden via afstemming op zijn resonantiefrequentie.",
    "Interatomic ether: Keely stelde dat tussen atomen een medium bestaat dat energie bevat — parallel aan het moderne kwantumvacuüm.",
    "Resonantie-levitatie: Demonstreerde voor journalisten en investeerders hoe zware metalen blokken konden worden opgeheven via muzikale resonantie.",
    "Disintegratie van graniet: Beweerde graniet te kunnen desintegreren met geluidsfrequenties — modern parallel: ultrasoon breken van steen is bewezen.",
]:
    doc.add_paragraph(f"  • {item}", style='List Bullet')

h3(doc, "De controverse")
body(doc, "Na Keely's dood in 1898 vond zijn huisbaas een elaborate ondergrondse tunnel met hydraulische apparatuur — mogelijk de aandrijving achter zijn 'ethermotoren'. Echter: sommige van zijn demonstraties, inclusief vloeistofdesintegratie, zijn nooit volledig verklaard. Modern ultrasoon onderzoek bevestigt enkele van zijn beweringen over resonantie en materiaalbreking. De waarheid ligt waarschijnlijk ergens tussen geniaal experimentator en theatraal bedrieger.")

divider(doc)

blueprint_box(doc, "KEELY'S SYMPATHETIC TRANSMITTER — Resonantie-principe (experimenteel)", [
    "  KEELY'S SYMPATHETIC VIBRATION PRINCIPE",
    "",
    "  GEDOCUMENTEERDE APPARAATSTRUCTUUR:",
    "",
    "  ┌──────────────────────────────────────────────────┐",
    "  │  RESONANTORKAMER (metalen cilinder)              │",
    "  │                                                  │",
    "  │  ┌────────┐  ┌────────┐  ┌────────┐            │",
    "  │  │STEM-   │  │STEM-   │  │STEM-   │            │",
    "  │  │VORK A  │  │VORK B  │  │VORK C  │            │",
    "  │  │(base)  │  │(5e)    │  │(octaaf)│            │",
    "  │  └────────┘  └────────┘  └────────┘            │",
    "  │       │           │           │                 │",
    "  │       └───────────┴───────────┘                 │",
    "  │                   │                             │",
    "  │              ACCOORD-RESONANTIE                 │",
    "  └───────────────────┬──────────────────────────────┘",
    "                      │",
    "               WATER- OF GLAZEN KOL",
    "               (resonantiemedium)",
    "",
    "  MODERN WERKEND EQUIVALENT:",
    "  Chladni figures: zand op metaalplaat toont resonantiepatronen",
    "  Cymatic patronen: water/zand in bakje toont frequentiestructuur",
    "",
    "  DEMO-EXPERIMENT (Chladni):",
    "  1. Metalen plaat Ø30cm op statief",
    "  2. Fijn zand of zout erop strooien",
    "  3. Strijkstok langs de rand (of sine tone generator + speaker eronder)",
    "  4. Zand vormt EXACTE geometrische patronen per frequentie",
    "  5. Dit is de zichtbare manifestatie van Keely's resonantieprincipe",
])

materials_box(doc, [
    "Metalen plaat (staal of messing) 30x30cm, 1-2mm dik",
    "Statief of houder van hout (niet-resonerend)",
    "Fijn kwartszand of keukenzout",
    "Cello- of altvioolstrijkstok (voor rand-excitatie)",
    "Alternatief: sine tone generator (app op telefoon) + luidspreker",
    "Frequentiemeter of tuner-app",
])

steps_box(doc, [
    "Bevestig metalen plaat horizontaal op statief (in het midden, zo min mogelijk opslaand).",
    "Strooi dun laagje fijn zand over het oppervlak.",
    "Houd plaat vast op een knooppunt en strijk met strijkstok langs de rand.",
    "Observeer hoe zand zich rangschikt in geometrische figuren — Chladni patronen.",
    "Varieer aanrakingspunt en strijkpositie voor andere patronen.",
    "Documenteer elke frequentie en zijn patroon — dit zijn 'staande golven' in materie.",
    "Geavanceerd: gebruik frequentiegenerator (audio) onder de plaat voor precise frequenties.",
])

doc.add_page_break()

# 12 — Antoine Béchamp
h2(doc, "  12    Antoine Béchamp — Het Terrein is Alles")
h3(doc, "Wie was hij?")
body(doc, "Antoine Béchamp (1816–1908) was een Franse chemicus en bioloog, tijdgenoot en rivaal van Louis Pasteur. Waar Pasteur's 'kiemtheorie' stelde dat specifieke externe ziekteverwekkers ziekte veroorzaken, stelde Béchamp dat het interne terrein — de biochemische omgeving van het lichaam — bepaalt of een organisme ziek wordt. Zijn theorie van de Microzym — kleine levende deeltjes in alle biologische materie — was radicaal en werd volledig verdrongen door het Pasteuriaanse paradigma dat onze gehele moderne geneeskunde aanstuurde.")

h3(doc, "Zijn ontdekkingen")
for item in [
    "Microzym-theorie: Kleine deeltjes (microzymas) in alle cellen kunnen transformeren in bacteriën afhankelijk van de biochemische omgeving.",
    "Pleomorfisme: Micro-organismen veranderen van vorm en functie afhankelijk van het biochemische milieu — niet alleen van externe kiemen.",
    "Terrein-theorie: Niet de kiem maar de biochemische conditie van het lichaam bepaalt of ziekte zich ontwikkelt.",
    "Fermentatie-onderzoek: Béchamp bewees dat fermentatie een biologisch, niet chemisch proces is — Pasteur nam deze ontdekking over zonder credit.",
]:
    doc.add_paragraph(f"  • {item}", style='List Bullet')

h3(doc, "Hoe werd het onderdrukt?")
body(doc, "Pasteur had betere politieke connecties en was een betere communicator. De kiemtheorie paste beter bij de farmaceutische industrie die specifieke middelen tegen specifieke kiemen kon verkopen. Béchamp's terrein-theorie vereiste een holistische, leefstijlgerichte aanpak — minder commercieel. Ironisch genoeg zei Pasteur op zijn sterfbed: 'Bernard had gelijk. De kiem is niets, het terrein is alles.' — Hoewel de authenticiteit van deze uitspraak wordt betwist.")
body(doc, "Moderne immunologie, het microbioom-onderzoek en epigenetica bewegen sterk in de richting van Béchamp's terrein-denken — zonder hem ooit expliciet te noemen.")

divider(doc)

blueprint_box(doc, "BÉCHAMP TERREIN-PROTOCOL — Biochemisch Milieu Optimaliseren", [
    "  BÉCHAMP TERREIN-OPTIMALISATIE PROTOCOL",
    "",
    "  KERNPRINCIPE: Optimaliseer het interne milieu →",
    "  pathogenen kunnen geen vaste voet krijgen",
    "",
    "  1. ZUUR-BASE BALANS:",
    "     Bloed pH optimaal: 7.35 - 7.45 (licht alkalisch)",
    "     Meting: urinestrook pH als indicator (niet exact maar indicatief)",
    "     Alkaliserend: groenten, citroensap (alkaliserend in lichaam!),",
    "                   mineraalwater, korte vasten",
    "     Verzurend: suiker, ultrabewerkt voedsel, chronische stress",
    "",
    "  2. OXIDATIEVE STRESS VERLAGEN:",
    "     Antioxidanten: vitamine C (1-3g/dag), glutathion precursors (NAC)",
    "     Polyfenolen: curcumine, quercetine, resveratrol",
    "     Zonlicht: vitamine D3 productie (terreinfactor, bewezen)",
    "",
    "  3. MICROBIOOM (Béchamp avant la lettre):",
    "     Gefermenteerd voedsel: kefir, kimchi, yoghurt",
    "     Prebiotica: inuline, psyllium, diverse plantenvezel",
    "     Antibiotica: enkel bij bewezen noodzaak (vernietigt terrein)",
    "",
    "  4. MINERAALDICHTHEID:",
    "     Magnesium: 400mg/dag (glycinaat voor opname)",
    "     Zink: 15-25mg/dag (immuunfunctie)",
    "     Selenium: 200mcg/dag (schildklier + immuun)",
    "     Jodium: 150-500mcg/dag (uit zeewier of supplement)",
    "",
    "  BÉCHAMP vs PASTEUR:",
    "  Pasteur: Kiem → Ziekte  (behandel de kiem)",
    "  Béchamp: Terrein → Kiem heeft kans  (behandel het terrein)",
    "  Modern: BEIDE zijn juist (host-pathogen interactie)",
])

doc.add_page_break()

# 13 — Nassim Haramein
h2(doc, "  13    Nassim Haramein — De Eenheid van Ruimte en Materie")
h3(doc, "Wie was hij?")
body(doc, "Nassim Haramein (1962–heden) is een Zwitserse theoretisch fysicus en onderzoeker die zonder academische achtergrond een unified field theory heeft ontwikkeld — de Resonance Science Foundation theorie, ook bekend als de Haramein-Rauscher metriek. Zijn kernstelling: het proton is een zwart gat, en de ruimte tussen materie is niet leeg maar een structuur van vacuümenergie (de 'planckse rooster') die alle materie verbindt.")
body(doc, "Zijn meest controversiële maar ook meest geciteerde paper — 'Quantum Gravity and the Holographic Mass' (2012) — werd gepubliceerd in een peer-reviewed tijdschrift en berekende de protonenstraal met verbluffende nauwkeurigheid vanuit eerste principes.")

h3(doc, "Zijn ontdekkingen")
for item in [
    "Proton als zwart gat: De informatie-inhoud van een proton, berekend als holoscherm, klopt met zijn massa — peer-reviewed bewijs.",
    "Connected Universe: Alles is verbonden via het kwantumvacuüm — de ruimte is het medium, niet de leegte.",
    "Torque and Coriolis in spin networks: Rotatie is fundamenteel aan alle schalen van het universum — van kwark tot supergalaxiscluster.",
    "64 tetrahedron grid: Geometrisch model voor het kwantumvacuüm dat experimenteel wordt getest.",
]:
    doc.add_paragraph(f"  • {item}", style='List Bullet')

h3(doc, "De controverse")
body(doc, "Haramein heeft geen formele universitaire graad en zijn theorieën zijn omstreden in de mainstream. Echter: zijn 2012 proton-massa berekening is niet weerlegd, en zijn voorspelde protonenstraal lag dichter bij experimentele waarden dan het standaardmodel-equivalent van die tijd. Hij vertegenwoordigt het nieuwe type buitenstaander-wetenschapper die peer-review-publicatie gebruikt terwijl hij buiten de institutionele structuur opereert.")

divider(doc)

blueprint_box(doc, "64 TETRAHEDRON GRID — Haramein's Kwantumvacuüm Geometrie", [
    "  HARAMEIN'S 64 TETRAHEDRON GRID (Isotrope Vector Matrix)",
    "",
    "  BASISSTRUCTUUR:",
    "  - 64 tetraëders gerangschikt in kuboctahedron-patroon",
    "  - Elke tetraëder deelt zijden met buren",
    "  - Resultaat: volledig gebalanceerde vector-matrix",
    "  - Grondpatroon: 8 Stella octangula's",
    "",
    "  VERBINDING MET BUCKMINSTER FULLER:",
    "  IVM (Isotrope Vector Matrix) = Fullers 'octet truss'",
    "  Dezelfde geometrie: sterkste structuur bij laagste materiaalgewicht",
    "  Gebruikt in: geodetische koepels, koolstofnanobuizen (C60)",
    "  Haramein: dit is de geometrie van het kwantumvacuüm zelf",
    "",
    "  MERKABA / DAVIDSTER VERBAND:",
    "  2 tetraëders in tegengestelde orientatie = Merkaba",
    "  64 tetrahedron grid = 32 Merkaba-paren",
    "  Oud-Egyptisch en Hebreeuws symbool = dezelfde geometrie",
    "",
    "  BOUW EEN 64 TETRAHEDRON MODEL:",
    "  Benodigdheden:",
    "  - 96 stokjes gelijke lengte (bamboe sate-stokjes werken)",
    "  - 24 verbindingsknooppunten (ballen van klei of 3D-print)",
    "  - Verdeel in 8 groepen van 12 stokjes per stella octangula",
    "",
    "  HARAMEIN'S PROTON BEREKENING (vereenvoudigd):",
    "  Holografisch principe: Info ~ Oppervlak / l_planck²",
    "  Massa proton = planckse massa × √(holografisch aantal)",
    "  Uitkomst: r_proton = 0.841 fm",
    "  CODATA 2018 experimenteel: 0.8414 fm  ← match!",
])

doc.add_page_break()

# ══════════════════════ OVERZICHT BLUEPRINTS ══════════════════════════════
h1(doc, "BIJLAGE A — OVERZICHT VAN ALLE BLUEPRINTS IN DIT BOEK")
doc.add_paragraph()

tbl = doc.add_table(rows=1, cols=4)
tbl.style = 'Table Grid'
header_cells = tbl.rows[0].cells
for cell, text in zip(header_cells, ['#', 'Blueprint', 'Uitvinder', 'Veiligheid']):
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    set_cell_bg(cell, '1A213E')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

blueprints_data = [
    ('1', 'Tesla Spoel (Spark Gap)', 'Nikola Tesla', '⚠️ HOOGSPANNING'),
    ('2', 'Wardenclyffe Transmissie Principe', 'Nikola Tesla', '⚠️ RF-straling'),
    ('3', 'Rife Frequentiegenerator (Arduino)', 'R.R. Rife', '⚠️ Geen medisch gebruik'),
    ('4', 'Rembrandt Impasto Laagopbouw', 'Rembrandt van Rijn', '✅ Veilig'),
    ('5', 'Semmelweis Chloorwas Protocol', 'I. Semmelweis', '✅ Veilig'),
    ('6', 'Orgon Accumulator', 'Wilhelm Reich', '⚡ Neutraal'),
    ('7', 'Hildegard Kruidentherapie', 'Hildegard van Bingen', '✅ Veilig'),
    ('8', 'Schauberger Vortex Pijp', 'Viktor Schauberger', '✅ Veilig'),
    ('9', 'Russell Octave Wave Tabel', 'Walter Russell', '✅ Theoretisch'),
    ('10', 'Biefeld-Brown Lifter', 'T.T. Brown', '⚠️ HOOGSPANNING 30kV'),
    ('11', 'Multiple Wave Oscillator', 'G. Lakhovsky', '⚠️ Hoogspanning'),
    ('12', 'Keely Chladni Demo', 'J.W. Keely', '✅ Veilig'),
    ('13', 'Béchamp Terrein Protocol', 'A. Béchamp', '✅ Veilig'),
    ('14', 'Haramein 64-Tetrahedron Grid', 'N. Haramein', '✅ Veilig'),
]

for row_data in blueprints_data:
    row = tbl.add_row()
    for cell, text in zip(row.cells, row_data):
        cell.text = text

doc.add_page_break()

# ══════════════════════ VERGELIJKINGSTABEL UITVINDERS ═════════════════════
h1(doc, "BIJLAGE B — VERGELIJKING: 13 GENIALE UITVINDERS VAN VERBORGEN KENNIS")
doc.add_paragraph()

tbl2 = doc.add_table(rows=1, cols=5)
tbl2.style = 'Table Grid'
header2 = tbl2.rows[0].cells
for cell, text in zip(header2, ['Uitvinder', 'Periode', 'Veld', 'Ontdekking', 'Status heden']):
    cell.text = text
    cell.paragraphs[0].runs[0].bold = True
    set_cell_bg(cell, '0F3460')
    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

inventors = [
    ('Tesla', '1856–1943', 'Elektriciteit', 'AC-stroom, draadloze energie', 'Erkend (AC), deels verborgen'),
    ('Rembrandt', '1606–1669', 'Kunst', 'Laagopbouw, lichtperceptie', 'Deels gereconstrueerd'),
    ('Rife', '1888–1971', 'Biofysica', 'Mortal Oscillatory Rate', 'PEMF erkend, MOR omstreden'),
    ('Semmelweis', '1818–1865', 'Geneeskunde', 'Handhygiëne = levensreddend', 'Volledig erkend (1880s)'),
    ('Reich', '1897–1957', 'Psychologie', 'Karakterpantsering, orgon', 'Somatiek erkend, orgon omstreden'),
    ('Hildegard', '1098–1179', 'Polymath', 'Kruidengeneeskunde, muziek', 'Heiligverklaard 2012'),
    ('Schauberger', '1885–1958', 'Hydrologie', 'Vortex water, implosie', 'Deels erkend in hydrologie'),
    ('Walter Russell', '1871–1963', 'Kosmologie', 'Uitgebreid periodiek systeem', 'Marginaal, voorspellingen kloppen'),
    ('T.T. Brown', '1905–1985', 'Fysica', 'Biefeld-Brown effect', 'Deels erkend, geclassificeerd'),
    ('Lakhovsky', '1869–1942', 'Biofysica', 'Multiple Wave Oscillator', 'Omstreden, onreplicated'),
    ('Keely', '1827–1898', 'Resonantie', 'Sympathische vibratiefysica', 'Deels bedrog, deels erkend'),
    ('Béchamp', '1816–1908', 'Biologie', 'Terrein-theorie, microzyms', 'Microbioom-revolutie bevestigt'),
    ('Haramein', '1962–heden', 'Kwantumfysica', 'Proton als zwart gat', 'Peer-reviewed, omstreden'),
]

for row_data in inventors:
    row = tbl2.add_row()
    for cell, text in zip(row.cells, row_data):
        cell.text = text

doc.add_page_break()

# ══════════════════════ CONCLUSIE ══════════════════════════════════════════
h2(doc, "  CONCLUSIE    Wat betekent dit voor jou?")

h3(doc, "De rode draad")
body(doc, "Al deze figuren deelden iets: ze zagen wat anderen niet zagen, of wilden niet zien. Ze werkten aan de rand van het acceptabele. Ze betaalden een prijs. En hun werk overleefde — hetzij in erkende vorm, hetzij als onzichtbaar fundament onder wat nu vanzelfsprekend heet.")
body(doc, "De vraag is niet 'wie heeft er gelijk gekregen'. De vraag is: welke ontdekkingen worden vandaag onderdrukt die over 50 jaar als vanzelfsprekend gelden? En welke rol speel jij — als denker, als bouwer, als mens?")

doc.add_paragraph()
h3(doc, "Zeven principes voor kritisch denken")
principes = [
    ("01", "Vraag altijd: wie heeft er baat bij dat dit niet bekend wordt?"),
    ("02", "Onderscheid wetenschappelijk bewijs van institutionele consensus."),
    ("03", "Documenteer je werk — voor de toekomst, niet alleen voor nu."),
    ("04", "Begrijp de financiële structuren achter kennis en informatie."),
    ("05", "Houd moed vast wanneer je gelijk hebt maar niemand luistert."),
    ("06", "Bouw netwerken die niet afhankelijk zijn van één geldstroom of autoriteit."),
    ("07", "Kennis is een mensenrecht (UVRM Art. 27) — deel wat je weet."),
]
for nr, tekst in principes:
    p = doc.add_paragraph()
    p.add_run(f"{nr}  ").bold = True
    p.add_run(tekst)

doc.add_paragraph()
quote(doc, "The present is theirs; the future, for which I really worked, is mine.", "Nikola Tesla")

doc.add_paragraph()
body(doc, "Dit document is onderdeel van het Aetherion Vectoris ecosysteem — gebouwd op de ethische handshake 0x416D69676F (Amigo) en verankerd in de Universele Verklaring van de Rechten van de Mens.")
body(doc, "Editie 2 — Juni 2026 | 13 uitvinders | 14 blueprints | Nabouwinstructies inbegrepen")

# ═══ OPSLAAN ═══════════════════════════════════════════════════════════════
output_path = "Verborgen_Kennis_Studieboek.docx"
doc.save(output_path)
print(f"[OK] Opgeslagen als {output_path}")
